from typing import TypedDict, Optional
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
import mysql.connector
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from docx import Document
import pandas as pd
import os
from docx.shared import Inches
from transformers import AutoTokenizer, AutoModel
import torch
import json
from decimal import Decimal
import sys
import time

# --- 🔎 상태 정의 ---
class ProposalState(TypedDict, total=False):
    brand_name: Optional[str]
    brand_info: Optional[str]
    client_needs: Optional[str]
    recent_issues: Optional[str]
    sales_status: Optional[str]
    recommended_media: Optional[str]
    previous_campaigns: Optional[str]
    proposal_text: Optional[str]
    proposal_file_path: Optional[str]
    media_info: list

AgentState = ProposalState

# --- .env 에서 OPENAI API 키 불러오기 ---
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

llm = ChatOpenAI(model="gpt-4o", openai_api_key=openai_api_key)

# --- HuggingFace 기반 임베딩 클래스 ---
class BERTSentenceEmbedding:
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModel.from_pretrained(model_name)

    def embed_documents(self, texts):
        return [self._embed(text) for text in texts]

    def embed_query(self, text):
        return self._embed(text)

    def _embed(self, text):
        inputs = self.tokenizer(text, return_tensors="pt", truncation=True, padding=True)
        with torch.no_grad():
            outputs = self.model(**inputs)
        cls_embedding = outputs.last_hidden_state[:, 0, :]
        return cls_embedding.squeeze(0).cpu().numpy()

embedding_function = BERTSentenceEmbedding()

# --- ✅ 기존 ChromaDB 불러오기 ---
vectorstore = Chroma(
    collection_name="campaign_media_chroma_hf",
    embedding_function=embedding_function,
    persist_directory="./chroma_db2"
)

print("기존 ChromaDB 로드 완료!", file=sys.stderr)

def db_query_tool(query: str):
    load_dotenv()
    conn = mysql.connector.connect(
        host=os.getenv("DB_HOST"),
        user=os.getenv("DB_USER"),
        password=os.getenv("DB_PASSWORD"),
        database=os.getenv("DB_NAME")
    )
    cursor = conn.cursor(dictionary=True)
    cursor.execute(query)
    results = cursor.fetchall()
    cursor.close()
    conn.close()
    return results

def web_search_tool(query: str) -> str:
    return f"[WEB SEARCH RESULT for: {query}]"

def vectordb_search_tool(query: str, vectorstore, top_k: int = 3) -> str:
    results = vectorstore.similarity_search(query, k=top_k)
    combined_results = []
    for doc in results:
        content = doc.page_content
        content_lines = [f"- {line.strip()}" for line in content.split(",")]
        content_formatted = "\n".join(content_lines)
        image_url = doc.metadata.get("execution_image_url", "")
        if image_url.startswith("/images/"):
            image_url = "./" + image_url.lstrip("/")
        elif image_url == "":
            image_url = "[이미지 없음]"
        content_with_image = f"{content_formatted}\n[이미지 보기]({image_url})"
        combined_results.append(content_with_image)
    return "\n\n---\n\n".join(combined_results)

def query_brand_and_sales_logs(brand_name: str):
    load_dotenv()
    conn = mysql.connector.connect(
        host=os.getenv("DB_HOST"),
        user=os.getenv("DB_USER"),
        password=os.getenv("DB_PASSWORD"),
        database=os.getenv("DB_NAME")
    )
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""SELECT * FROM brands WHERE brand_name = %s""", (brand_name,))
    brand_info = cursor.fetchone()

    if not brand_info:
        cursor.close()
        conn.close()
        return None, None

    brand_id = brand_info["brand_id"]
    cursor.execute("""SELECT * FROM sales_logs WHERE brand_id = %s ORDER BY contact_time DESC LIMIT 1""", (brand_id,))
    latest_sales_log = cursor.fetchone()
    cursor.close()
    conn.close()

    return brand_info, latest_sales_log

def analyze_brand_and_needs(state: ProposalState):
    brand_name = state["brand_name"]
    brand_info, latest_sales_log = query_brand_and_sales_logs(brand_name)
    if not brand_info:
        raise ValueError(f"브랜드 '{brand_name}' 정보를 찾을 수 없습니다.")

    client_needs = latest_sales_log["client_needs_summary"] if latest_sales_log else "최근 고객 요구사항 정보 없음"
    recent_issues = brand_info.get("recent_brand_issues") or "브랜드 이슈 정보 없음"
    sales_status = brand_info.get("sales_status") or "상태 정보 없음"

    return {
        **state,
        "brand_info": brand_info,
        "client_needs": client_needs,
        "recent_issues": recent_issues,
        "sales_status": sales_status
    }

def retrieve_previous_campaigns(state: ProposalState):
    client_needs = state.get("client_needs") or "옥외 광고 집행 사례"
    similar_cases = vectordb_search_tool(client_needs, vectorstore)
    return {**state, "previous_campaigns": similar_cases}

def recommend_media(state: ProposalState):
    client_needs = state.get("client_needs") or ""
    db_results = db_query_tool("SELECT * FROM medias WHERE quantity > 0;")
    if not db_results:
        raise ValueError("사용 가능한 매체 정보가 없습니다.")
    media_info = db_results

    media_json = json.dumps(db_results, ensure_ascii=False, default=lambda o: float(o) if isinstance(o, Decimal) else str(o))
    prompt = f"""
        당신은 옥외 광고 전문 대행사의 전략 기획자입니다.
        다음 브랜드의 고객 요구사항과 유사 집행 사례를 고려하여 가장 적합한 옥외 광고 매체 3가지를 추천해야 합니다.

        - 브랜드 고객 요구사항: {state.get('client_needs')}
        - 유사 집행 사례 요약: {state.get('previous_campaigns')}

        다음은 사용 가능한 매체 리스트입니다:
        {media_json}
    """
    recommendation = llm.invoke(prompt)
    recommendation_text = recommendation.content

    return {**state, "recommended_media": recommendation_text, "media_info": media_info}

def generate_proposal(state: ProposalState):
    import datetime
    import re
    from docx import Document
    from docx.shared import Inches
    from langchain.prompts import ChatPromptTemplate

    doc = Document()
    doc.add_heading(f"{state['brand_name']} 옥외 광고 제안서", level=1)
    doc.add_paragraph("") 

    # --- 1. 고객사 정보 ---
    doc.add_heading("1. 고객사 정보", level=2)
    brand_info = state["brand_info"]

    # brand_info가 dict 형태라면:
    if isinstance(brand_info, dict):
        for key, value in brand_info.items():
            doc.add_paragraph(f"- {key}: {value}")
    else:
        # CSV처럼 한 줄로 들어왔을 경우 → 키:값 형태로 줄바꿈
        lines = re.split(r",|\t", brand_info)
        for line in lines:
            if ':' in line:
                doc.add_paragraph(f"- {line.strip()}")
            elif line.strip():
                doc.add_paragraph(f"- {line.strip()}")

    # --- 2. 캠페인 목표 ---
    doc.add_paragraph("") 
    doc.add_heading("2. 캠페인 목표", level=2)
    client_needs = state["client_needs"]

    # 쉼표로 구분 → 문장별로 출력
    for item in re.split(r",|·|•", client_needs):
        if item.strip():
            doc.add_paragraph(f"- {item.strip()}")

    # --- 3. 유사 집행 사례 ---
    doc.add_heading("3. 유사 집행 사례", level=2)
    previous_campaigns = state["previous_campaigns"]

    cases = previous_campaigns.split("\n\n---\n\n")
    filtered_cases = []
    for idx, case in enumerate(cases, start=1):
        if "사례 3" in case:
            continue  # 사례 3 제외

        # 이미지 URL 추출
        image_url = None
        image_match = re.search(r"\[이미지 보기\]\((.*?)\)", case)
        if image_match:
            image_url = image_match.group(1).strip()

        # 텍스트에서 [이미지 보기] 부분 삭제
        case_text = re.sub(r"\[이미지 보기\]\(.*?\)", "", case).strip()
        filtered_cases.append((f"- 사례 {idx}\n{case_text}", image_url))

    # ✅ 3행 2열짜리 표 생성 (행 = 사례 개수, 열 = 2)
    table = doc.add_table(rows=len(filtered_cases), cols=2)
    table.style = 'Table Grid'

    for row_idx, (case_text, image_url) in enumerate(filtered_cases):
        # 왼쪽 셀: 사례 내용
        table.cell(row_idx, 0).text = case_text

        # 오른쪽 셀: 이미지 또는 텍스트
        cell_image = table.cell(row_idx, 1)
        if image_url and image_url.endswith((".jpg", ".png")):
            try:
                run = cell_image.paragraphs[0].add_run()
                run.add_picture(image_url, width=Inches(2.5))
            except Exception as e:
                cell_image.text = f"이미지 삽입 실패 ({e})"
        else:
            cell_image.text = "이미지 없음"

    # --- 4. 추천매체 및 집행계획 ---
    doc.add_paragraph("") 
    doc.add_heading("4. 추천매체 및 집행계획", level=2)

    recommended_media = state["recommended_media"]
    if hasattr(recommended_media, "content"):
        recommended_media_text = recommended_media.content
    else:
        recommended_media_text = recommended_media

    media_rows = state.get("media_info", [])
    media_image_map = {row["media_name"].strip().lower(): row["image_day_url"] for row in media_rows}

    # ⭐ "1.", "2.", "3." 으로 시작하는 추천 매체 블록만 추출
    lines = recommended_media_text.split("\n")
    media_blocks = []
    current_block = ""

    for line in lines:
        clean_line = re.sub(r"\*+", "", line.strip())
        if re.match(r"^\d+\.", clean_line):  # "1.", "2.", "3." 매체 번호로 시작
            if current_block:
                media_blocks.append(current_block.strip())
            current_block = clean_line
        elif current_block:
            current_block += "\n" + clean_line
    # 마지막 블록 추가
    if current_block:
        media_blocks.append(current_block.strip())

    # ⭐ 상위 3개 매체만 사용
    selected_media_blocks = media_blocks[:3]

    # ✅ 3행 2열짜리 표 생성
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    for idx, block in enumerate(selected_media_blocks):
        # 매체명 추출
        name_match = re.search(r"매체명\s*[:：]\s*(.+)", block)
        if name_match:
            media_name = name_match.group(1).strip()
        else:
            # 🔍 예: "매체 id: 14 - n.square 강남대로"
            name_match = re.search(r"매체 id\s*[:：]?\s*\d+\s*-\s*(.+)", block)
            if name_match:
                media_name = name_match.group(1).strip()
            else:
                # 🔍 예: "1. n.square 강남대로"
                name_match = re.search(r"^\d+\.\s*(.+)", block)
                media_name = name_match.group(1).strip() if name_match else "추천 매체"

        media_name_clean = media_name.replace("*", "").strip().lower()
        media_name_clean = re.sub(r"\(media id: \d+\)", "", media_name_clean).strip()
        image_url = media_image_map.get(media_name_clean)
        
        with open("media_debug_log.txt", "a", encoding="utf-8") as log_file:
                log_file.write(f"[{idx}] 매체명: {media_name_clean} → 이미지 경로: {image_url}\n")

        if image_url and image_url.startswith("/images/"):
            image_url = "./images/" + image_url[len("/images/"):]

        # 왼쪽 셀: 매체 설명
        table.cell(idx, 0).text = block

        # 오른쪽 셀: 매체 이미지
        cell_image = table.cell(idx, 1)
        # 이미지 삽입
        if image_url and image_url.endswith((".jpg", ".png")):
            try:
                run = cell_image.paragraphs[0].add_run()
                run.add_picture(image_url, width=Inches(3))
            except Exception as e:
                cell_image.text = f"이미지 삽입 실패 ({e})"
        else:
            cell_image.text = "이미지 없음"

    # 혹시 3개보다 적으면 나머지 비워놓기
    for idx in range(len(selected_media_blocks), 3):
        table.cell(idx, 0).text = ""
        table.cell(idx, 1).text = ""

    # --- 5. 결론 (LLM 사용) ---
    doc.add_paragraph("") 
    doc.add_heading("5. 결론", level=2)

    # LLM 프롬프트
    prompt = ChatPromptTemplate.from_template("""
    브랜드명: {brand_name}
    캠페인 목표: {client_needs}
    추천 매체: {recommended_media}

    위 정보를 바탕으로 제안서의 마무리 결론 부분을 작성하세요.
    """)
    chain = prompt | llm
    conclusion = chain.invoke(state).content

    doc.add_paragraph(conclusion)

    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"{state['brand_name']}_제안서_{now}.docx"
    doc.save(file_name)

    return {**state, "proposal_text": conclusion, "proposal_file_path": file_name}

# --- 🔗 그래프 구성 ---
graph = StateGraph(ProposalState)
graph.add_node("AnalyzeBrandAndNeeds", analyze_brand_and_needs)
graph.add_node("RecommendMedia", recommend_media)
graph.add_node("RetrievePreviousCampaigns", retrieve_previous_campaigns)
graph.add_node("GenerateProposal", generate_proposal)

graph.set_entry_point("AnalyzeBrandAndNeeds")
graph.add_edge("AnalyzeBrandAndNeeds", "RetrievePreviousCampaigns")
graph.add_edge("RetrievePreviousCampaigns", "RecommendMedia")
graph.add_edge("RecommendMedia", "GenerateProposal")
graph.set_finish_point("GenerateProposal")

proposal_graph = graph.compile()

# --- 🚀 실행 ---
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("--brand", required=True, help="브랜드명 (예: 유니클로코리아)")
    args = parser.parse_args()

    initial_state = {
        "brand_name": args.brand
    }

    final_state = proposal_graph.invoke(initial_state)

    print("최종 제안서:\n", file=sys.stderr)
    print(final_state["proposal_text"], file=sys.stderr)
    print(f"제안서 Word 파일 경로: {final_state['proposal_file_path']}", file=sys.stderr)

    result = {
        "success": True,
        "brand": initial_state["brand_name"],
        "file_path": final_state["proposal_file_path"],
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    }
    print(json.dumps(result))