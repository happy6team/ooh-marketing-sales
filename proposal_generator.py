from typing import TypedDict, Optional
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
import mysql.connector
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from docx import Document
import pandas as pd
import os
from transformers import AutoTokenizer, AutoModel
import torch
import json
from decimal import Decimal
import sys
import time
from sqlalchemy import text
import asyncio

from db import AsyncSessionLocal

# --- 🔎 상태 정의 ---
class ProposalState(TypedDict, total=False):
    brand_name: Optional[str]
    brand_info: Optional[str]
    client_needs: Optional[str]
    recent_issues: Optional[str]
    sales_status: Optional[str]
    recommended_media: Optional[str]
    previous_campaigns: Optional[str]
    proposal_text: Optional[str]
    proposal_file_path: Optional[str]
    media_info: list

AgentState = ProposalState

# --- .env 에서 OPENAI API 키 불러오기 ---
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

llm = ChatOpenAI(model="gpt-4o", openai_api_key=openai_api_key)

# --- HuggingFace 기반 임베딩 클래스 ---
class BERTSentenceEmbedding:
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModel.from_pretrained(model_name)

    def embed_documents(self, texts):
        return [self._embed(text) for text in texts]

    def embed_query(self, text):
        return self._embed(text)

    def _embed(self, text):
        inputs = self.tokenizer(text, return_tensors="pt", truncation=True, padding=True)
        with torch.no_grad():
            outputs = self.model(**inputs)
        cls_embedding = outputs.last_hidden_state[:, 0, :]
        return cls_embedding.squeeze(0).cpu().numpy()

embedding_function = BERTSentenceEmbedding()

# --- ✅ 기존 ChromaDB 불러오기 ---
vectorstore = Chroma(
    collection_name="campaign_media_chroma_hf",
    embedding_function=embedding_function,
    persist_directory="../chroma_db2"
)

print("기존 ChromaDB 로드 완료!", file=sys.stderr)

async def db_query_tool_async(query: str, params: tuple = None):
    async with AsyncSessionLocal() as session:
        try:
            async with session.begin():
                stmt = text(query)
                result = await session.execute(stmt, params or ())
                return [dict(row._mapping) for row in result.fetchall()]
        except Exception as e:
            return []

def web_search_tool(query: str) -> str:
    return f"[WEB SEARCH RESULT for: {query}]"

def vectordb_search_tool(query: str, vectorstore, top_k: int = 3) -> str:
    results = vectorstore.similarity_search(query, k=top_k)
    combined_results = []
    for doc in results:
        content = doc.page_content
        content_lines = [f"- {line.strip()}" for line in content.split(",")]
        content_formatted = "\n".join(content_lines)
        image_url = doc.metadata.get("execution_image_url", "")
        if image_url.startswith("/images/"):
            image_url = "../" + image_url.lstrip("/")
        elif image_url == "":
            image_url = "[이미지 없음]"
        content_with_image = f"{content_formatted}\n[이미지 보기]({image_url})"
        combined_results.append(content_with_image)
    return "\n\n---\n\n".join(combined_results)

async def query_brand_and_sales_logs(brand_name: str) -> tuple[dict, Optional[dict]]:
    brand_info_list = await db_query_tool_async("SELECT * FROM brand WHERE brand_name = %s", (brand_name,))
    if not brand_info_list:
        return {}, None

    brand_info = brand_info_list[0]
    brand_id = brand_info["brand_id"]
    sales_logs = await db_query_tool_async(
        "SELECT * FROM sales_log WHERE brand_id = %s ORDER BY contact_time DESC LIMIT 1", (brand_id,)
    )
    return brand_info, sales_logs[0] if sales_logs else None

async def analyze_brand_and_needs(state: ProposalState):
    brand_name = state["brand_name"]
    brand_info, latest_sales_log = await query_brand_and_sales_logs(brand_name)
    if not brand_info:
        raise ValueError(f"브랜드 '{brand_name}' 정보를 찾을 수 없습니다.")

    client_needs = latest_sales_log["client_needs_summary"] if latest_sales_log else "최근 고객 요구사항 정보 없음"
    recent_issues = brand_info.get("recent_brand_issues") or "브랜드 이슈 정보 없음"
    sales_status = brand_info.get("sales_status") or "상태 정보 없음"

    return {
        **state,
        "brand_info": brand_info,
        "client_needs": client_needs,
        "recent_issues": recent_issues,
        "sales_status": sales_status
    }

def retrieve_previous_campaigns(state: ProposalState):
    client_needs = state.get("client_needs") or "옥외 광고 집행 사례"
    similar_cases = vectordb_search_tool(client_needs, vectorstore)
    return {**state, "previous_campaigns": similar_cases}

async def recommend_media(state: ProposalState):
    client_needs = state.get("client_needs") or ""
    db_results = await db_query_tool_async("SELECT * FROM media WHERE quantity > 0;")
    if not db_results:
        raise ValueError("사용 가능한 매체 정보가 없습니다.")
    media_info = db_results

    media_json = json.dumps(db_results, ensure_ascii=False, default=lambda o: float(o) if isinstance(o, Decimal) else str(o))
    prompt = f"""
        당신은 옥외 광고 전문 대행사의 전략 기획자입니다.
        다음 브랜드의 고객 요구사항과 유사 집행 사례를 고려하여 가장 적합한 옥외 광고 매체 3가지를 추천해야 합니다.

        - 브랜드 고객 요구사항: {state.get('client_needs')}
        - 유사 집행 사례 요약: {state.get('previous_campaigns')}

        다음은 사용 가능한 매체 리스트입니다:
        {media_json}
    """
    recommendation = await llm.invoke(prompt)
    recommendation_text = recommendation.content

    return {**state, "recommended_media": recommendation_text, "media_info": media_info}

def generate_proposal(state: ProposalState):
    import datetime
    import re
    doc = Document()
    doc.add_heading(f"{state['brand_name']} 옥외 광고 제안서", level=1)

    doc.add_heading("1. 고객사 정보", level=2)
    brand_info = state["brand_info"]
    if isinstance(brand_info, dict):
        for key, value in brand_info.items():
            doc.add_paragraph(f"- {key}: {value}")

    doc.add_heading("2. 캠페인 목표", level=2)
    client_needs = state["client_needs"]
    for item in re.split(r",|·|•", client_needs):
        if item.strip():
            doc.add_paragraph(f"- {item.strip()}")

    doc.add_heading("3. 유사 집행 사례", level=2)
    previous_campaigns = state["previous_campaigns"]
    cases = previous_campaigns.split("\n\n---\n\n")
    for idx, case in enumerate(cases, 1):
        doc.add_paragraph(f"- 사례 {idx}: {case}")

    doc.add_heading("4. 추천매체 및 집행계획", level=2)
    doc.add_paragraph(state["recommended_media"])

    doc.add_heading("5. 결론", level=2)
    prompt = ChatPromptTemplate.from_template("""
    브랜드명: {brand_name}
    캠페인 목표: {client_needs}
    추천 매체: {recommended_media}

    위 정보를 바탕으로 제안서의 마무리 결론 부분을 작성하세요.
    """)
    chain = prompt | llm
    conclusion = chain.invoke(state).content
    doc.add_paragraph(conclusion)

    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"{state['brand_name']}_제안서_{now}.docx"
    doc.save(file_name)

    return {**state, "proposal_text": conclusion, "proposal_file_path": file_name}

# --- 🔗 그래프 구성 ---
graph = StateGraph(ProposalState)
graph.add_node("AnalyzeBrandAndNeeds", analyze_brand_and_needs)
graph.add_node("RecommendMedia", recommend_media)
graph.add_node("RetrievePreviousCampaigns", retrieve_previous_campaigns)
graph.add_node("GenerateProposal", generate_proposal)

graph.set_entry_point("AnalyzeBrandAndNeeds")
graph.add_edge("AnalyzeBrandAndNeeds", "RetrievePreviousCampaigns")
graph.add_edge("RetrievePreviousCampaigns", "RecommendMedia")
graph.add_edge("RecommendMedia", "GenerateProposal")
graph.set_finish_point("GenerateProposal")

proposal_graph = graph.compile()

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--brand", required=True)
    args = parser.parse_args()

    try:
        initial_state = {"brand_name": args.brand}
        final_state = asyncio.run(proposal_graph.ainvoke(initial_state))
        
        print("✅ 제안서 생성 완료", file=sys.stderr)
        print(final_state["proposal_text"], file=sys.stderr)
        print(f"📄 저장 경로: {final_state['proposal_file_path']}", file=sys.stderr)

        result = {
            "success": True,
            "brand": args.brand,
            "file_path": final_state["proposal_file_path"],
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        print(json.dumps(result))
    except Exception as e:
        result = {"success": False, "error": str(e)}
        print(json.dumps(result))
