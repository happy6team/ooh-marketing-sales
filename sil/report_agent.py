from typing import TypedDict, Optional
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
import mysql.connector
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from docx import Document
import pandas as pd
import os
from docx.shared import Inches
from transformers import AutoTokenizer, AutoModel
import torch
import json
from decimal import Decimal


# --- 🔎 상태 정의 ---
class ProposalState(TypedDict, total=False):
    brand_name: Optional[str]
    brand_info: Optional[str]
    client_needs: Optional[str]
    recent_issues: Optional[str]
    sales_status: Optional[str]
    recommended_media: Optional[str]
    previous_campaigns: Optional[str]
    proposal_text: Optional[str]
    proposal_file_path: Optional[str]
    media_info: list

AgentState = ProposalState

# --- .env 에서 OPENAI API 키 불러오기 ---
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

llm = ChatOpenAI(model="gpt-4o", openai_api_key=openai_api_key)

# --- HuggingFace 기반 임베딩 클래스 ---
class BERTSentenceEmbedding:
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModel.from_pretrained(model_name)

    def embed_documents(self, texts):
        return [self._embed(text) for text in texts]

    def embed_query(self, text):
        return self._embed(text)

    def _embed(self, text):
        inputs = self.tokenizer(text, return_tensors="pt", truncation=True, padding=True)
        with torch.no_grad():
            outputs = self.model(**inputs)
        cls_embedding = outputs.last_hidden_state[:, 0, :]  # (batch_size, hidden_size)
        return cls_embedding.squeeze(0).cpu().numpy()

# --- embedding_function 인스턴스 생성 ---
embedding_function = BERTSentenceEmbedding()

# --- 🗄️ Tool 구성 (여기서는 예시 Stub 형태로 Tool 설정) ---
# 실제 서비스에서는 아래 Tool들을 LangChain Toolkit으로 구현

def db_query_tool(query: str):
    import mysql.connector
    from dotenv import load_dotenv
    import os

    load_dotenv()  # .env 파일 로드

    host = os.getenv("DB_HOST")
    user = os.getenv("DB_USER")
    password = os.getenv("DB_PASSWORD")
    database = os.getenv("DB_NAME")

    conn = mysql.connector.connect(
        host=host,
        user=user,
        password=password,
        database=database
    )
    cursor = conn.cursor(dictionary=True)
    cursor.execute(query)
    results = cursor.fetchall()
    cursor.close()
    conn.close()
    return results


def web_search_tool(query: str) -> str:
    return f"[WEB SEARCH RESULT for: {query}]"

def vectordb_search_tool(query: str, collection_name: str, top_k: int = 3) -> str:
    vectorstore = Chroma(
        collection_name=collection_name,
        embedding_function=embedding_function
    )
    results = vectorstore.similarity_search(query, k=top_k)

    combined_results = []
    for doc in results:
        content = doc.page_content

        content_lines = [f"- {line.strip()}" for line in content.split(",")]
        content_formatted = "\n".join(content_lines)

        image_url = doc.metadata.get("execution_image_url", "")
        if image_url.startswith("/images/"):
            image_url = "../" + image_url.lstrip("/")
        elif image_url == "":
            image_url = "[이미지 없음]"

        content_with_image = f"{content_formatted}\n[이미지 보기]({image_url})"
        combined_results.append(content_with_image)

    return "\n\n---\n\n".join(combined_results)

# --- ChromaDB에 campaign_media 데이터 올리기 (최초 1회) ---
csv_path = "../data/data_sample/campaign_media.csv"
df = pd.read_csv(csv_path)

def row_to_text(row):
    return (
        f"캠페인 ID: {row['campaign_id']}, "
        f"매체 ID: {row['media_id']}, "
        f"시작일: {row['start_date']}, "
        f"종료일: {row['end_date']}, "
        f"구좌 수: {row['slot_count']}, "
        f"집행 가격: {row['executed_price']}, "
        f"진행 상태: {row['campaign_media_status']}"
    )

texts = []
metadatas = []

for idx, row in df.iterrows():
    texts.append(row_to_text(row))
    metadatas.append({"execution_image_url": row["execution_image_url"]})

# --- 문서 분할 ---
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
docs = splitter.create_documents(texts, metadatas=metadatas)

# --- Chroma 컬렉션에 저장 (HuggingFace 임베딩 사용) ---
vectorstore = Chroma.from_documents(
    documents=docs,
    embedding=embedding_function,  # ✅ embedding_function으로 수정
    collection_name="campaign_media_chroma_hf"  # ✅ 새 컬렉션 이름
)

print(f"✅ HuggingFace 임베딩으로 {len(docs)}개 문서 저장 완료!")

def query_brand_and_sales_logs(brand_name: str):
    load_dotenv()
    conn = mysql.connector.connect(
        host=os.getenv("DB_HOST"),
        user=os.getenv("DB_USER"),
        password=os.getenv("DB_PASSWORD"),
        database=os.getenv("DB_NAME")
    )
    cursor = conn.cursor(dictionary=True)

    # 1️⃣ 브랜드 정보 가져오기
    cursor.execute("""
        SELECT * FROM brand WHERE brand_name = %s
    """, (brand_name,))
    brand_info = cursor.fetchone()

    if not brand_info:
        cursor.close()
        conn.close()
        return None, None

    brand_id = brand_info["brand_id"]

    # 2️⃣ 최신 sales_log 정보 가져오기 (brand_id 기준)
    cursor.execute("""
        SELECT * FROM sales_log 
        WHERE brand_id = %s 
        ORDER BY contact_time DESC 
        LIMIT 1
    """, (brand_id,))
    latest_sales_log = cursor.fetchone()

    cursor.close()
    conn.close()

    return brand_info, latest_sales_log

def decimal_default(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    raise TypeError

# --- Node 1: 브랜드 정보 + 고객 요구사항 분석 ---
def analyze_brand_and_needs(state: ProposalState):
    brand_name = state["brand_name"]

    brand_info, latest_sales_log = query_brand_and_sales_logs(brand_name)

    if not brand_info:
        raise ValueError(f"브랜드 '{brand_name}' 정보를 찾을 수 없습니다.")

    # 최신 고객 요구사항 요약 가져오기
    client_needs = latest_sales_log["client_needs_summary"] if latest_sales_log else "최근 고객 요구사항 정보 없음"

    # 브랜드 이슈 및 상태 정보 추가 (LLM 프롬프트에 도움됨)
    recent_issues = brand_info.get("recent_brand_issues") or "브랜드 이슈 정보 없음"
    sales_status = brand_info.get("sales_status") or "상태 정보 없음"

    return {
        **state,
        "brand_info": brand_info,
        "client_needs": client_needs,
        "recent_issues": recent_issues,
        "sales_status": sales_status
    }

# --- Node 2: 유사 집행 사례 조회 사진 정보 먼저---
def retrieve_previous_campaigns(state: ProposalState):
    client_needs = state.get("client_needs") or "옥외 광고 집행 사례"

    collection_name = "campaign_media_chroma_hf"  # 벡터스토어 컬렉션 이름
    similar_cases = vectordb_search_tool(client_needs, collection_name)

    return {**state, "previous_campaigns": similar_cases}

# --- Node 3: 매체 추천 및 매칭 유사 집행 사례와 기존 매체 통합해서 MZ 패키지 매체 여러개 ---
# media, 웹검색?!
def recommend_media(state: ProposalState):
    client_needs = state.get("client_needs") or ""
    db_results = db_query_tool("SELECT * FROM media WHERE quantity > 0;")

    if not db_results or len(db_results) == 0:
        raise ValueError("사용 가능한 매체 정보가 없습니다.")

    media_info = db_results  # 이걸 state에 넘김

    media_json = json.dumps(db_results, ensure_ascii=False, default=lambda o: float(o) if isinstance(o, Decimal) else str(o))

    prompt = f"""
        당신은 옥외 광고 전문 대행사의 전략 기획자입니다.

        다음 브랜드의 고객 요구사항과 유사 집행 사례를 고려하여 가장 적합한 옥외 광고 매체 3가지를 추천해야 합니다.

        - 브랜드 고객 요구사항: {state.get('client_needs')}
        - 유사 집행 사례 요약: {state.get('previous_campaigns')}

        다음은 사용 가능한 매체 리스트입니다:
        {media_json}

        각각 다음 기준으로 추천하십시오:
        1. 대행사 입장에서 가장 전략적으로 추천하는 매체 (효과와 인지도가 높음)
        2. 가격적으로 저렴하면서 효과적인 매체
        3. 기타 추천할 만한 매체 한 가지

        각 후보 매체에 대해 다음 정보를 포함하십시오:
        - 매체명: {{매체명}}
        - 예상 집행 가격: {{금액}} 원
        - 예상 집행 기간: {{기간}}
        - 예상 노출 빈도: {{노출 빈도}}
        - 추천 이유: {{이유}}

        **매체명: 이라는 키워드를 꼭 포함하고 각 매체는 줄바꿈하여 구분하세요.**
        """

    recommendation = llm.invoke(prompt)
    recommendation_text = recommendation.content  # 문자열로 변환
    
    return {
        **state,
        "recommended_media": recommendation_text,
        "media_info": media_info
    }

# --- Node 4: 제안서 생성 (Word 파일 포함) ---
def generate_proposal(state: ProposalState):
    import datetime
    import re
    from docx import Document
    from docx.shared import Inches
    from langchain.prompts import ChatPromptTemplate

    doc = Document()
    doc.add_heading(f"{state['brand_name']} 옥외 광고 제안서", level=1)
    doc.add_paragraph("") 

    # --- 1. 고객사 정보 ---
    doc.add_heading("1. 고객사 정보", level=2)
    brand_info = state["brand_info"]

    # brand_info가 dict 형태라면:
    if isinstance(brand_info, dict):
        for key, value in brand_info.items():
            doc.add_paragraph(f"- {key}: {value}")
    else:
        # CSV처럼 한 줄로 들어왔을 경우 → 키:값 형태로 줄바꿈
        lines = re.split(r",|\t", brand_info)
        for line in lines:
            if ':' in line:
                doc.add_paragraph(f"- {line.strip()}")
            elif line.strip():
                doc.add_paragraph(f"- {line.strip()}")

    # --- 2. 캠페인 목표 ---
    doc.add_paragraph("") 
    doc.add_heading("2. 캠페인 목표", level=2)
    client_needs = state["client_needs"]

    # 쉼표로 구분 → 문장별로 출력
    for item in re.split(r",|·|•", client_needs):
        if item.strip():
            doc.add_paragraph(f"- {item.strip()}")

    # --- 3. 유사 집행 사례 ---
    doc.add_heading("3. 유사 집행 사례", level=2)
    previous_campaigns = state["previous_campaigns"]

    cases = previous_campaigns.split("\n\n---\n\n")
    filtered_cases = []
    for idx, case in enumerate(cases, start=1):
        if "사례 3" in case:
            continue  # 사례 3 제외

        # 이미지 URL 추출
        image_url = None
        image_match = re.search(r"\[이미지 보기\]\((.*?)\)", case)
        if image_match:
            image_url = image_match.group(1).strip()

        # 텍스트에서 [이미지 보기] 부분 삭제
        case_text = re.sub(r"\[이미지 보기\]\(.*?\)", "", case).strip()
        filtered_cases.append((f"- 사례 {idx}\n{case_text}", image_url))

    # ✅ 3행 2열짜리 표 생성 (행 = 사례 개수, 열 = 2)
    table = doc.add_table(rows=len(filtered_cases), cols=2)
    table.style = 'Table Grid'

    for row_idx, (case_text, image_url) in enumerate(filtered_cases):
        # 왼쪽 셀: 사례 내용
        table.cell(row_idx, 0).text = case_text

        # 오른쪽 셀: 이미지 또는 텍스트
        cell_image = table.cell(row_idx, 1)
        if image_url and image_url.endswith((".jpg", ".png")):
            try:
                run = cell_image.paragraphs[0].add_run()
                run.add_picture(image_url, width=Inches(2.5))
            except Exception as e:
                cell_image.text = f"이미지 삽입 실패 ({e})"
        else:
            cell_image.text = "이미지 없음"

    # --- 4. 추천매체 및 집행계획 ---
    doc.add_paragraph("") 
    doc.add_heading("4. 추천매체 및 집행계획", level=2)

    recommended_media = state["recommended_media"]
    if hasattr(recommended_media, "content"):
        recommended_media_text = recommended_media.content
    else:
        recommended_media_text = recommended_media

    media_rows = state.get("media_info", [])
    media_image_map = {row["media_name"].strip().lower(): row["image_day_url"] for row in media_rows}

    media_blocks = []
    current_block = ""
    for line in recommended_media_text.split("\n"):
        # ⭐ 제목의 "**" 제거 (block 추가 전 처리)
        clean_line = re.sub(r"\*+", "", line.strip())
        if re.match(r"^\d+\.", clean_line):
            if current_block:
                media_blocks.append(current_block.strip())
            current_block = clean_line
        elif clean_line:
            current_block += "\n" + clean_line
    if current_block:
        media_blocks.append(current_block.strip())

    table = doc.add_table(rows=len(media_blocks), cols=2)
    table.style = 'Table Grid'

    for idx, block in enumerate(media_blocks):
        # 매체명 추출
        name_match = re.search(r"매체명\s*[:：]\s*(.+)", block)
        if name_match:
            media_name = name_match.group(1).strip()
        else:
            name_match = re.search(r"^\d+\.\s*(.+)", block)
            media_name = name_match.group(1).strip() if name_match else "추천 매체"

        media_name_clean = media_name.replace("*", "").strip().lower()
        image_url = media_image_map.get(media_name_clean)
        if image_url and image_url.startswith("/images/"):
            image_url = "../images/" + image_url[len("/images/"):]
        # 왼쪽: 설명
        table.cell(idx, 0).text = block
        # 오른쪽: 이미지
        cell_image = table.cell(idx, 1)
        if image_url and image_url.endswith((".jpg", ".png")):
            try:
                run = cell_image.paragraphs[0].add_run()
                run.add_picture(image_url, width=Inches(3))
            except Exception as e:
                cell_image.text = f"이미지 삽입 실패 ({e})"
        else:
            cell_image.text = "이미지 없음"

    # --- 5. 결론 (LLM 사용) ---
    doc.add_paragraph("") 
    doc.add_heading("5. 결론", level=2)

    # LLM 프롬프트
    prompt = ChatPromptTemplate.from_template("""
    브랜드명: {brand_name}
    캠페인 목표: {client_needs}
    추천 매체: {recommended_media}

    위 정보를 바탕으로 캠페인의 결론 부분을 작성하세요.
    """)
    chain = prompt | llm
    conclusion = chain.invoke(state).content

    doc.add_paragraph(conclusion)

    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"{state['brand_name']}_제안서_{now}.docx"
    doc.save(file_name)

    return {**state, "proposal_text": conclusion, "proposal_file_path": file_name}

# --- 🔗 그래프 구성 ---
graph = StateGraph(ProposalState)

graph.add_node("AnalyzeBrandAndNeeds", analyze_brand_and_needs)
graph.add_node("RecommendMedia", recommend_media)
graph.add_node("RetrievePreviousCampaigns", retrieve_previous_campaigns)
graph.add_node("GenerateProposal", generate_proposal)

graph.set_entry_point("AnalyzeBrandAndNeeds")
graph.add_edge("AnalyzeBrandAndNeeds", "RetrievePreviousCampaigns")
graph.add_edge("RetrievePreviousCampaigns", "RecommendMedia")
graph.add_edge("RecommendMedia", "GenerateProposal")
graph.set_finish_point("GenerateProposal")

proposal_graph = graph.compile()

# --- 🚀 실행 예시 ---
initial_state = {
    "brand_name": "유니클로코리아",
}

final_state = proposal_graph.invoke(initial_state)

print("✅ 최종 제안서:\n")
print(final_state["proposal_text"])
print(f"📄 제안서 Word 파일 경로: {final_state['proposal_file_path']}")